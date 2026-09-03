#!/usr/bin/env python3
"""Convert a Markdown file to a .docx file.

Usage:
    python md2docx.py <input.md> <output.docx>

Supports ATX headings (#..######), paragraphs, unordered lists (-/*),
ordered lists (1.), and pipe tables. Inline formatting (bold/italic/code)
is not parsed; content is written as plain runs. This is a lossy, plain-text
writer intended for simple documents, not round-tripping styled files.
"""
import re
import sys

from docx import Document


def parse_table(lines):
    """Parse a pipe-table block (header, separator, rows) into list of rows."""
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln.startswith("|"):
            break
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
    # Drop the separator row (---) if present as second row.
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        del rows[1]
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    src, dst = sys.argv[1], sys.argv[2]
    with open(src, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # Table block: starts with | and next line is a separator.
        if stripped.startswith("|") and i + 1 < n and re.fullmatch(
            r"\s*\|?[\s:|-]+\|?\s*", lines[i + 1]
        ) and "---" in lines[i + 1]:
            j = i
            while j < n and lines[j].strip().startswith("|"):
                j += 1
            rows = parse_table(lines[i:j])
            add_table(doc, rows)
            i = j
            continue
        # Headings.
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue
        # Unordered list.
        if re.match(r"^[-*]\s+", stripped):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue
        # Ordered list.
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            doc.add_paragraph(m.group(1), style="List Number")
            i += 1
            continue
        # Plain paragraph.
        doc.add_paragraph(stripped)
        i += 1

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main()
